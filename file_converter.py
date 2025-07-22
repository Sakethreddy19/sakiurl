import streamlit as st
from pdf2docx import Converter
from docx import Document
import pandas as pd
import os
import pdfkit
import comtypes.client  # Windows only
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
import time

# Configure wkhtmltopdf
pdfkit_config = pdfkit.configuration(wkhtmltopdf='C:/Program Files/wkhtmltopdf/bin/wkhtmltopdf.exe')

# ---------- Conversion Functions ----------

def convert_pdf_to_docx(input_pdf, output_docx):
    try:
        cv = Converter(input_pdf)
        cv.convert(output_docx)
        cv.close()
        return True, "PDF converted to DOCX."
    except Exception as e:
        return False, f"PDF→DOCX failed: {str(e)}"

def convert_docx_to_pdf(input_docx, output_pdf):
    try:
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(os.path.abspath(input_docx))
        doc.SaveAs(os.path.abspath(output_pdf), FileFormat=17)
        doc.Close()
        word.Quit()
        return True, "DOCX converted to PDF."
    except Exception as e:
        return False, f"DOCX→PDF failed: {str(e)}"

def convert_excel_to_pdf(input_excel, output_pdf):
    try:
        df = pd.read_excel(input_excel)
        html_file = "temp_excel.html"
        df.to_html(html_file)
        pdfkit.from_file(html_file, output_pdf, configuration=pdfkit_config)
        os.remove(html_file)
        return True, "Excel converted to PDF."
    except Exception as e:
        return False, f"Excel→PDF failed: {str(e)}"

def convert_excel_to_docx(input_excel, output_docx):
    try:
        df = pd.read_excel(input_excel)
        doc = Document()
        doc.add_heading('Excel Data', 0)
        table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])

        for j, column in enumerate(df.columns):
            table.cell(0, j).text = str(column)
        for i, row in df.iterrows():
            for j, val in enumerate(row):
                table.cell(i+1, j).text = str(val)

        doc.save(output_docx)
        return True, "Excel converted to DOCX."
    except Exception as e:
        return False, f"Excel→DOCX failed: {str(e)}"

# ---------- Web Scraping Function ----------

def scrape_website(url):
    try:
        chrome_options = Options()
        chrome_options.add_argument("--headless")
        chrome_options.add_argument('--disable-gpu')
        chrome_options.add_argument("--no-sandbox")
        chrome_options.add_argument("--disable-dev-shm-usage")

        driver = webdriver.Chrome(options=chrome_options)
        driver.get(url)
        time.sleep(5)

        title = driver.title

        try:
            meta_desc = driver.find_element(By.XPATH, '//meta[@name="description"]').get_attribute("content")
        except:
            meta_desc = "No description found."

        paragraphs = driver.find_elements(By.TAG_NAME, 'p')
        content = "\n".join([p.text for p in paragraphs if p.text.strip() != ""])

        driver.quit()

        full_text = f"Title: {title}\n\nDescription: {meta_desc}\n\nContent:\n{content}"
        return full_text
    except Exception as e:
        return f"Scraping failed: {str(e)}"

# ---------- Streamlit App ----------

st.title("🔁 File Converter + 🌐 Web Scraper")

tab1, tab2 = st.tabs(["📁 File Conversion", "🌐 Web Scraping"])

# -------- Tab 1: File Converter --------
with tab1:
    uploaded_file = st.file_uploader("Upload your file", type=["pdf", "docx", "xlsx"])
    conversion_type = st.selectbox("Choose conversion", [
        "PDF to DOCX", "DOCX to PDF", "Excel to PDF", "Excel to DOCX"
    ])

    if st.button("Convert File"):
        if not uploaded_file:
            st.warning("Please upload a file.")
        else:
            ext = uploaded_file.name.rsplit(".", 1)[-1].lower()
            input_path = f"temp_input.{ext}"
            with open(input_path, "wb") as f:
                f.write(uploaded_file.read())

            base = uploaded_file.name.rsplit(".", 1)[0]
            output_file = ""

            if conversion_type == "PDF to DOCX":
                output_file = f"{base}_converted.docx"
                success, msg = convert_pdf_to_docx(input_path, output_file)

            elif conversion_type == "DOCX to PDF":
                output_file = f"{base}_converted.pdf"
                success, msg = convert_docx_to_pdf(input_path, output_file)

            elif conversion_type == "Excel to PDF":
                output_file = f"{base}_converted.pdf"
                success, msg = convert_excel_to_pdf(input_path, output_file)

            elif conversion_type == "Excel to DOCX":
                output_file = f"{base}_converted.docx"
                success, msg = convert_excel_to_docx(input_path, output_file)

            if success:
                st.success(msg)
                with open(output_file, "rb") as f:
                    st.download_button("⬇️ Download File", f, file_name=output_file)
                os.remove(output_file)
            else:
                st.error(msg)

            os.remove(input_path)

# -------- Tab 2: Web Scraper --------
with tab2:
    url = st.text_input("Enter Article URL (e.g., Times of India, Hindustan Times)")
    format = st.radio("Export scraped content as:", ["DOCX", "PDF"])

    if st.button("Scrape and Export"):
        if not url.strip():
            st.warning("Please enter a valid URL.")
        else:
            scraped_text = scrape_website(url)
            if "Scraping failed" in scraped_text:
                st.error(scraped_text)
            else:
                st.text_area("Preview", scraped_text[:1000], height=300)

                if format == "DOCX":
                    doc = Document()
                    doc.add_heading("Scraped Content", 0)
                    doc.add_paragraph(scraped_text)
                    doc.save("scraped.docx")
                    with open("scraped.docx", "rb") as f:
                        st.download_button("⬇️ Download DOCX", f, file_name="scraped.docx")
                    os.remove("scraped.docx")

                elif format == "PDF":
                    with open("scraped.html", "w", encoding="utf-8") as f:
                        f.write(f"<html><body><h1>Scraped Content</h1><pre>{scraped_text}</pre></body></html>")
                    pdfkit.from_file("scraped.html", "scraped.pdf", configuration=pdfkit_config)
                    with open("scraped.pdf", "rb") as f:
                        st.download_button("⬇️ Download PDF", f, file_name="scraped.pdf")
                    os.remove("scraped.html")
                    os.remove("scraped.pdf")
